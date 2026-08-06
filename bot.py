import os
import asyncio
import logging
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import CommandStart
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from openai import AsyncOpenAI
import docx

BOT_TOKEN = "8695864774:AAH8V_XNxTlvXP89EFk6MnF7bAdKNxCxS5A"
OPENAI_API_KEY = "dawlet1215..."

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()
ai_client = AsyncOpenAI(api_key=OPENAI_API_KEY)

class WorkOrder(StatesGroup):
    choosing_language = State()
    waiting_for_single_form = State()
    waiting_for_text_translation = State()
    waiting_for_file_translation = State()

def get_lang_keyboard():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="Qaraqalpaqsha 🚩", callback_data="lang_kaa"),
                InlineKeyboardButton(text="O'zbekcha 🇺🇿", callback_data="lang_uz"),
            ],
            [
                InlineKeyboardButton(text="Қазақша 🇰🇿", callback_data="lang_kk"),
                InlineKeyboardButton(text="Русский 🇷🇺", callback_data="lang_ru"),
            ],
            [
                InlineKeyboardButton(text="English 🇬🇧", callback_data="lang_en"),
            ]
        ]
    )

def get_main_keyboard_kaa():
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="📄 Referat", callback_data="btn_referat"),
                InlineKeyboardButton(text="🎬 Slayd", callback_data="btn_slayd"),
            ],
            [
                InlineKeyboardButton(text="📝 Maqala", callback_data="btn_maqala"),
                InlineKeyboardButton(text="💡 Tezis", callback_data="btn_tezis"),
            ],
            [
                InlineKeyboardButton(text="📘 Kurs jumısı", callback_data="btn_kurs"),
                InlineKeyboardButton(text="📋 Test", callback_data="btn_test"),
            ],
            [
                InlineKeyboardButton(text="✍️ Tekst Audarmashı", callback_data="btn_text_trans"),
                InlineKeyboardButton(text="📁 Fayl Audarmashı", callback_data="btn_file_trans"),
            ],
            [
                InlineKeyboardButton(text="☎️ Baylanıs", callback_data="btn_baylanys"),
            ]
        ]
    )

def get_main_keyboard_other(lang):
    if lang == "uz":
        labels = ["📄 Referat", "🎬 Slayd", "📝 Maqola", "💡 Tezis", "📘 Kurs ishi", "📋 Test", "✍️ Matn Tarjima", "📁 Fayl Tarjima", "☎️ Aloqa"]
    elif lang == "kk":
        labels = ["📄 Реферат", "🎬 Слайд", "📝 Мақала", "💡 Тезис", "📘 Курстық жұмыс", "📋 Тест", "✍️ Мәтін Аударма", "📁 Файл Аударма", "☎️ Байланыс"]
    elif lang == "ru":
        labels = ["📄 Реферат", "🎬 Слайд", "📝 Статья", "💡 Тезис", "📘 Курсовая работа", "📋 Тест", "✍️ Перевод Текста", "📁 Перевод Файла", "☎️ Контакты"]
    else:
        labels = ["📄 Paper", "🎬 Slide", "📝 Article", "💡 Thesis", "📘 Coursework", "📋 Test", "✍️ Text Translation", "📁 File Translation", "☎️ Contact"]

    return InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text=labels[0], callback_data="btn_referat"),
                InlineKeyboardButton(text=labels[1], callback_data="btn_slayd"),
            ],
            [
                InlineKeyboardButton(text=labels[2], callback_data="btn_maqala"),
                InlineKeyboardButton(text=labels[3], callback_data="btn_tezis"),
            ],
            [
                InlineKeyboardButton(text=labels[4], callback_data="btn_kurs"),
                InlineKeyboardButton(text=labels[5], callback_data="btn_test"),
            ],
            [
                InlineKeyboardButton(text=labels[6], callback_data="btn_text_trans"),
                InlineKeyboardButton(text=labels[7], callback_data="btn_file_trans"),
            ],
            [
                InlineKeyboardButton(text=labels[8], callback_data="btn_baylanys"),
            ]
        ]
    )

@dp.message(CommandStart())
async def start_handler(message: types.Message, state: FSMContext):
    await state.clear()
    welcome_text = (
        f"Assalawma áleykum, {message.from_user.first_name}! 👋\n\n"
        "🤖 **Bilim AI** botına xosh keldińiz!\n"
        "Iltimas, ózińizge qolaylı tildi saylań / Выберите язык:"
    )
    await message.answer(welcome_text, reply_markup=get_lang_keyboard())
    await state.set_state(WorkOrder.choosing_language)

@dp.callback_query(F.data.in_(["lang_kaa", "lang_uz", "lang_kk", "lang_ru", "lang_en"]))
async def set_language(callback: types.CallbackQuery, state: FSMContext):
    lang = callback.data.replace("lang_", "")
    await state.update_data(user_lang=lang)
    
    msg_dict = {
        "kaa": "🚩 Qaraqalpaq tili saylandı! Tómendegi xızmet bólimlerinen birin saylań:",
        "uz": "🇺🇿 O'zbek tili tanlandi! Quyidagi tugmalardan birini tanlang:",
        "kk": "🇰🇿 Қазақ тілі таңдалды! Төмендегі түймелердің бірін таңдаңыз:",
        "ru": "🇷🇺 Выбран русский язык! Выберите одну из кнопок ниже:",
        "en": "🇬🇧 English language selected! Choose one of the buttons below:"
    }
    kb = get_main_keyboard_kaa() if lang == "kaa" else get_main_keyboard_other(lang)
    await callback.message.answer(msg_dict[lang], reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data == "btn_baylanys")
async def contact_handler(callback: types.CallbackQuery, state: FSMContext):
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    contact_text = "☎️ **Baylanıs / Contact**\n\n📩 Telegram: @sizning_username\n📞 Telefon: +998 XX XXX XX XX"
    kb = get_main_keyboard_kaa() if lang == "kaa" else get_main_keyboard_other(lang)
    await callback.message.answer(contact_text, parse_mode="Markdown", reply_markup=kb)
    await callback.answer()

@dp.callback_query(F.data == "btn_text_trans")
async def text_trans_start(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(WorkOrder.waiting_for_text_translation)
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    
    msg = {
        "kaa": "✍️ **Tekst Audarmashı bólimi**\n\nAudarma etiliwi kerek bolǵan tekstti jiberiń:",
        "uz": "✍️ **Matn Tarjima bo'limi**\n\nTarjima qilinishi kerak bo'lgan matnni yuboring:",
        "kk": "✍️ **Мәтін Аударма бөлімі**\n\nАударылуы керек мәтінді жіберіңіз:",
        "ru": "✍️ **Переводчик Текста**\n\nОтправьте текст для перевода:",
        "en": "✍️ **Text Translator**\n\nSend the text you want to translate:"
    }
    await callback.message.answer(msg[lang])
    await callback.answer()

@dp.message(WorkOrder.waiting_for_text_translation)
async def process_text_translation(message: types.Message, state: FSMContext):
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    
    wait_msg = await message.answer("⏳ Audarılmaqta, sabır etiń...")
    kb = get_main_keyboard_kaa() if lang == "kaa" else get_main_keyboard_other(lang)
    
    try:
        sys_prompt = f"Translate the given text strictly to target language '{lang}' (for 'kaa', use Karakalpak Latin script). Preserve original formatting."
        response = await ai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": sys_prompt},
                {"role": "user", "content": message.text}
            ]
        )
        await wait_msg.delete()
        await message.answer(response.choices[0].message.content, reply_markup=kb)
    except Exception:
        await wait_msg.delete()
        await message.answer("Keshiriń, qátelik júz berdi.", reply_markup=kb)
    await state.clear()

@dp.callback_query(F.data == "btn_file_trans")
async def file_trans_start(callback: types.CallbackQuery, state: FSMContext):
    await state.set_state(WorkOrder.waiting_for_file_translation)
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    
    msg = {
        "kaa": "📁 **Fayl Audarmashı bólimi**\n\nFayldagı barlıq tekstler, kesteler hám sýretler astındagı maǵlıwmatlar tolıq audarıladı.\nAudarma etetuǵın faylıńızdı (DOCX, TXT) jiberiń:",
        "uz": "📁 **Fayl Tarjima bo'limi**\n\nFayldagi barcha matnlar, jadvallar va rasmlar to'liq tarjima qilinadi (DOCX, TXT):",
        "kk": "📁 **Файл Аударма бөлімі**\n\nФайлды жіберіңіз (DOCX, TXT):",
        "ru": "📁 **Переводчик Файлов**\n\nОтправьте файл. Все тексты, таблицы и описания графиков будут полностью переведены (DOCX, TXT):",
        "en": "📁 **File Translator**\n\nSend your file (DOCX, TXT). Text, tables, and image texts will be translated fully:"
    }
    await callback.message.answer(msg[lang])
    await callback.answer()

@dp.message(WorkOrder.waiting_for_file_translation, F.document)
async def process_file_translation(message: types.Message, state: FSMContext):
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    kb = get_main_keyboard_kaa() if lang == "kaa" else get_main_keyboard_other(lang)
    
    wait_msg = await message.answer("⏳ Fayl júklep alınıp, tolıq maǵlıwmatlar (tekst, keste, grafik maǵlıwmatları) audarılmaqta...")
    file_id = message.document.file_id
    file_name = message.document.file_name
    
    file = await bot.get_file(file_id)
    download_path = f"downloads/{file_name}"
    os.makedirs("downloads", exist_ok=True)
    await bot.download_file(file.file_path, download_path)
    
    try:
        if file_name.endswith(".docx"):
            doc = docx.Document(download_path)
            
            for p in doc.paragraphs:
                if p.text.strip():
                    res = await ai_client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": f"Translate to target language '{lang}' (use Karakalpak Latin for 'kaa'). Keep original tone."},
                            {"role": "user", "content": p.text}
                        ]
                    )
                    p.text = res.choices[0].message.content
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            res = await ai_client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": f"Translate strictly to target language '{lang}' (Karakalpak Latin for 'kaa')."},
                                    {"role": "user", "content": cell.text}
                                ]
                            )
                            cell.text = res.choices[0].message.content

            out_path = f"downloads/Audarılǵan_{file_name}"
            doc.save(out_path)
            
            await wait_msg.delete()
            await message.answer_document(types.FSInputFile(out_path), caption="✅ Faylıńızdaǵı barlıq tekst hám kesteler tolıq audarıldı!", reply_markup=kb)
        else:
            await wait_msg.delete()
            await message.answer("Házirshe tolıq format saqlaw ushın DOCX faylları qabıllanadı.", reply_markup=kb)
    except Exception as e:
        await wait_msg.delete()
        await message.answer("Keshiriń, fayldı qayta islewde qátelik júz berdi.", reply_markup=kb)
    
    await state.clear()

@dp.callback_query(F.data.in_(["btn_referat", "btn_slayd", "btn_maqala", "btn_tezis", "btn_kurs", "btn_test"]))
async def single_page_form_start(callback: types.CallbackQuery, state: FSMContext):
    service_names = {
        "btn_referat": "📄 Referat",
        "btn_slayd": "🎬 Slayd",
        "btn_maqala": "📝 Maqala",
        "btn_tezis": "💡 Tezis",
        "btn_kurs": "📘 Kurs jumısı (100 betge deyin)",
        "btn_test": "📋 Test"
    }
    chosen_service = service_names.get(callback.data, "Akademiyalıq jumıs")
    await state.update_data(chosen_service=chosen_service)
    await state.set_state(WorkOrder.waiting_for_single_form)
    
    form_text = (
        f"📋 **{chosen_service} ushın maǵlıwmatlardı bir jerde toltırıń:**\n\n"
        "Tómendegi shablondı kóshirip alıp, ózińizge kerekli maǵlıwmatlar menen toltırıp, botqa BIR XABAR etip jiberiń:\n\n"
        "```text\n"
        f"Xızmet: {chosen_service}\n"
        "Tema: \n"
        "Bet/Slayd sanı (Maksimum 100 bet / 30 slayd): \n"
        "Til: Qaraqalpaqsha (Latin)\n"
        "Slayd dizayn stili (Tek Slayd ushın): IT / Akademiyalıq / Minimalizm / Business\n"
        "Oqıw ornı hám Kafedra: \n"
        "Student (Atı-jóni, Kursı): \n"
        "Muǵallim (Atı-jóni, Dárejesi): \n"
        "Qosımsha: [☑️ Keste] [☑️ Foto/Diagramma] [☑️ Kod/Algoritm]\n"
        "```\n\n"
        "👇 *Shablondı kóshirip alıp, toltırıp jiberiń:*"
    )
    await callback.message.answer(form_text, parse_mode="Markdown")
    await callback.answer()

@dp.message(WorkOrder.waiting_for_single_form)
async def process_single_form_result(message: types.Message, state: FSMContext):
    user_data = await state.get_data()
    lang = user_data.get("user_lang", "kaa")
    kb = get_main_keyboard_kaa() if lang == "kaa" else get_main_keyboard_other(lang)
    
    wait_msg = await message.answer("⚡️ **AI iske tústi...**\n⏳ *Maǵlıwmatlar qayta islenip, Titul beti, grafikler, kesteler, foto-bloklar hám kodlar menen akademiyalıq jumıs tayarlanbaqta...*")
    
    system_prompt = (
        "Siz akademiyalıq jumıslardı tolıq tayarlawshı ekspert-AI siz. "
        "Paydalanıwshı bergen barlıq maǵlıwmatlar tiykarında Titul betin, Mazmunın (plan), "
        "Tiykarǵı teoriyalıq hám amalıy bólimlerdi, Kestelerdi, Foto/Diagramma orınların, "
        "Kod blokların hám Paydalanılǵan ádebiyatlar dizimin tayarlap beriń. "
        "JUWAPTI TEK SAP QARAQALPAQ TILINDE LATIN ALIPBESINDE JAZIŃ."
    )
    
    try:
        response = await ai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": message.text}
            ]
        )
        await wait_msg.delete()
        await message.answer(response.choices[0].message.content, reply_markup=kb)
    except Exception:
        await wait_msg.delete()
        await message.answer("Keshiriń, jumıstı tayarlawda qátelik júz berdi.", reply_markup=kb)
    
    await state.clear()

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    asyncio.run(main())
  
